from __future__ import annotations

from io import BytesIO

import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.documents.document_content import (
    DocumentBlock,
    DocumentContent,
)


BODY_FONT = "Noto Sans CJK KR"


STYLE_PROFILES = {
    "corporate_clean": {
        "title_size": 20,
        "heading_size": 12,
        "accent": "1F4E78",
        "label_fill": "EAF0F5",
        "header_fill": "D9E6F2",
        "callout_fill": "F3F6F9",
    },
    "formal_korean": {
        "title_size": 20,
        "heading_size": 12,
        "accent": "222222",
        "label_fill": "E7E6E6",
        "header_fill": "D9D9D9",
        "callout_fill": "F2F2F2",
    },
    "executive_report": {
        "title_size": 22,
        "heading_size": 13,
        "accent": "17365D",
        "label_fill": "DCE6F1",
        "header_fill": "B8CCE4",
        "callout_fill": "EAF2F8",
    },
    "modern_project": {
        "title_size": 21,
        "heading_size": 12,
        "accent": "244062",
        "label_fill": "EAF2F8",
        "header_fill": "D6E4F0",
        "callout_fill": "EEF4F8",
    },
    "compact_memo": {
        "title_size": 18,
        "heading_size": 11,
        "accent": "333333",
        "label_fill": "EEEEEE",
        "header_fill": "E2E2E2",
        "callout_fill": "F5F5F5",
    },
    "minimal": {
        "title_size": 18,
        "heading_size": 11,
        "accent": "333333",
        "label_fill": "F2F2F2",
        "header_fill": "EBEBEB",
        "callout_fill": "F7F7F7",
    },
}


def _profile(name: str) -> dict:
    return STYLE_PROFILES.get(
        name,
        STYLE_PROFILES["corporate_clean"],
    )


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"),
        BODY_FONT,
    )

    if size is not None:
        run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold

    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))

    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)

    shading.set(qn("w:fill"), fill)


def _add_heading(
    document: Document,
    text: str,
) -> None:
    text = text.strip()

    if not text:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _add_paragraph(
    document: Document,
    text: str,
) -> None:
    text = text.strip()

    if not text:
        return

    for line in text.splitlines():
        line = line.strip()

        if not line:
            continue

        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.15


def _add_key_value_table(
    document: Document,
    data: dict[str, str],
) -> None:
    items = [
        (key.strip(), value.strip())
        for key, value in data.items()
        if key.strip() and value.strip()
    ]

    if not items:
        return

    table = document.add_table(
        rows=0,
        cols=4,
    )
    table.style = "Table Grid"

    for index in range(0, len(items), 2):
        cells = table.add_row().cells

        key1, value1 = items[index]
        cells[0].text = key1
        cells[1].text = value1

        for run in cells[0].paragraphs[0].runs:
            run.bold = True

        if index + 1 < len(items):
            key2, value2 = items[index + 1]
            cells[2].text = key2
            cells[3].text = value2

            for run in cells[2].paragraphs[0].runs:
                run.bold = True


def _add_table(
    document: Document,
    rows: list[list[str]],
) -> None:
    cleaned_rows = [
        [str(cell).strip() for cell in row]
        for row in rows
        if row
    ]

    if not cleaned_rows:
        return

    column_count = max(
        len(row)
        for row in cleaned_rows
    )

    if column_count <= 0:
        return

    table = document.add_table(
        rows=0,
        cols=column_count,
    )
    table.style = "Table Grid"

    for row_index, row in enumerate(cleaned_rows):
        cells = table.add_row().cells

        for column_index in range(column_count):
            value = (
                row[column_index]
                if column_index < len(row)
                else ""
            )

            cells[column_index].text = value

            if row_index == 0:
                for run in cells[column_index].paragraphs[0].runs:
                    run.bold = True


def _add_bullet_list(
    document: Document,
    items: list[str],
) -> None:
    for item in items:
        text = item.strip()

        if not text:
            continue

        paragraph = document.add_paragraph(
            style="List Bullet",
        )
        paragraph.add_run(text)


def _add_numbered_list(
    document: Document,
    items: list[str],
) -> None:
    for item in items:
        text = item.strip()

        if not text:
            continue

        paragraph = document.add_paragraph(
            style="List Number",
        )
        paragraph.add_run(text)


def _add_callout(
    document: Document,
    block: DocumentBlock,
) -> None:
    text = block.content.strip()

    if not text and not block.title.strip():
        return

    table = document.add_table(
        rows=1,
        cols=1,
    )
    table.style = "Table Grid"

    cell = table.cell(0, 0)

    if block.title.strip():
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(
            block.title.strip()
        )
        run.bold = True

        if text:
            cell.add_paragraph(text)
    else:
        cell.text = text


def _add_signature(
    document: Document,
    block: DocumentBlock,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if block.title.strip():
        run = paragraph.add_run(
            block.title.strip()
        )
        run.bold = True

    if block.content.strip():
        if block.title.strip():
            paragraph.add_run("\n")

        paragraph.add_run(
            block.content.strip()
        )

    for key, value in block.data.items():
        if not key.strip() or not value.strip():
            continue

        paragraph.add_run(
            f"\n{key.strip()}: {value.strip()}"
        )


def _render_block(
    document: Document,
    block: DocumentBlock,
) -> None:
    if block.type == "page_break":
        document.add_page_break()
        return

    if block.type == "heading":
        _add_heading(
            document,
            block.content or block.title,
        )
        return

    if block.title.strip():
        _add_heading(
            document,
            block.title,
        )

    if block.type == "paragraph":
        _add_paragraph(
            document,
            block.content,
        )

    elif block.type == "key_value_table":
        _add_key_value_table(
            document,
            block.data,
        )

    elif block.type == "table":
        _add_table(
            document,
            block.rows,
        )

    elif block.type == "bullet_list":
        _add_bullet_list(
            document,
            block.items,
        )

    elif block.type == "numbered_list":
        _add_numbered_list(
            document,
            block.items,
        )

    elif block.type == "callout":
        _add_callout(
            document,
            block,
        )

    elif block.type == "signature":
        _add_signature(
            document,
            block,
        )


def _render_legacy_content(
    document: Document,
    doc: DocumentContent,
) -> None:
    if doc.body.strip():
        _add_paragraph(
            document,
            doc.body,
        )

    for index, section in enumerate(
        doc.sections,
        start=1,
    ):
        if section.title.strip():
            _add_heading(
                document,
                f"{index}. {section.title.strip()}",
            )

        if section.content.strip():
            _add_paragraph(
                document,
                section.content,
            )


def _set_paragraph_bottom_border(
    paragraph,
    color: str,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))

    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _apply_business_style(
    document: Document,
    doc: DocumentContent,
) -> None:
    profile = _profile(doc.style_profile)

    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        BODY_FONT,
    )

    paragraphs = document.paragraphs

    if paragraphs:
        title = paragraphs[0]
        title.paragraph_format.space_after = Pt(14)

        for run in title.runs:
            _set_run_font(
                run,
                size=profile["title_size"],
                bold=True,
                color=profile["accent"],
            )

        _set_paragraph_bottom_border(
            title,
            profile["accent"],
        )

    heading_number = 0

    for paragraph in paragraphs[1:]:
        if not paragraph.runs:
            continue

        is_heading = any(
            run.bold
            and run.font.size
            and 11 <= run.font.size.pt <= 13
            for run in paragraph.runs
        )

        if is_heading:
            heading_number += 1

            text = paragraph.text.strip()

            if (
                text
                and not re.match(
                    r"^\d+[\.\)]\s*",
                    text,
                )
            ):
                paragraph.runs[0].text = (
                    f"{heading_number}. "
                    f"{paragraph.runs[0].text}"
                )

            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_with_next = True

            for run in paragraph.runs:
                _set_run_font(
                    run,
                    size=profile["heading_size"],
                    bold=True,
                    color=profile["accent"],
                )

        else:
            paragraph.paragraph_format.line_spacing = 1.2

            for run in paragraph.runs:
                _set_run_font(
                    run,
                    size=10.5,
                )

    for table_index, table in enumerate(
        document.tables
    ):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        is_metadata = (
            bool(doc.metadata)
            and table_index == 0
            and len(table.columns) == 4
        )

        is_callout = (
            len(table.rows) == 1
            and len(table.columns) == 1
        )

        for row_index, row in enumerate(
            table.rows
        ):
            for column_index, cell in enumerate(
                row.cells
            ):
                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

                    for run in paragraph.runs:
                        _set_run_font(
                            run,
                            size=9.5,
                        )

                if (
                    is_metadata
                    and column_index in {0, 2}
                ):
                    _shade_cell(
                        cell,
                        profile["label_fill"],
                    )

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                elif (
                    not is_metadata
                    and not is_callout
                    and row_index == 0
                ):
                    _shade_cell(
                        cell,
                        profile["header_fill"],
                    )

                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                elif is_callout:
                    _shade_cell(
                        cell,
                        profile["callout_fill"],
                    )


def render_docx(
    doc: DocumentContent,
) -> bytes:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    title_run = title.add_run(
        doc.title.strip() or "문서"
    )
    title_run.bold = True
    title_run.font.size = Pt(18)

    if doc.metadata:
        _add_key_value_table(
            document,
            doc.metadata,
        )

        document.add_paragraph("")

    if doc.blocks:
        for block in doc.blocks:
            _render_block(
                document,
                block,
            )
    else:
        _render_legacy_content(
            document,
            doc,
        )

    _apply_business_style(
        document,
        doc,
    )

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()
