from __future__ import annotations

import re
from io import BytesIO

from docx import Document as DocxDocument

from app.services.documents.document_content import (
    DocumentBlock,
    DocumentContent,
)


def _normalize(text: str) -> str:
    text = re.sub(
        r"^\s*\d+\s*[.)]\s*",
        "",
        str(text or "").strip(),
    )
    return re.sub(r"\s+", " ", text).strip().lower()


def _block_text(block: DocumentBlock) -> str:
    if block.type == "paragraph":
        return block.content.strip()

    if block.type == "bullet_list":
        return "\n".join(
            f"• {item.strip()}"
            for item in block.items
            if item.strip()
        )

    if block.type == "numbered_list":
        return "\n".join(
            f"{index}. {item.strip()}"
            for index, item in enumerate(
                block.items,
                start=1,
            )
            if item.strip()
        )

    if block.type == "key_value_table":
        return "\n".join(
            f"{key.strip()}: {value.strip()}"
            for key, value in block.data.items()
            if key.strip() and value.strip()
        )

    if block.type == "table":
        return "\n".join(
            " | ".join(
                str(cell).strip()
                for cell in row
            )
            for row in block.rows
            if row
        )

    if block.type in {
        "callout",
        "signature",
    }:
        values = []

        if block.content.strip():
            values.append(block.content.strip())

        values.extend(
            f"{key.strip()}: {value.strip()}"
            for key, value in block.data.items()
            if key.strip() and value.strip()
        )

        return "\n".join(values)

    return block.content.strip()


def _collect_sections(
    doc: DocumentContent,
) -> dict[str, str]:
    sections: dict[str, list[str]] = {}

    current_heading: str | None = None

    for block in doc.blocks:
        if block.type == "heading":
            current_heading = (
                block.content
                or block.title
            ).strip()

            if current_heading:
                sections.setdefault(
                    current_heading,
                    [],
                )

            continue

        if not current_heading:
            continue

        text = _block_text(block)

        if text:
            sections[current_heading].append(text)

    return {
        title: "\n".join(parts).strip()
        for title, parts in sections.items()
        if parts
    }


def _replace_placeholders(
    text: str,
    values: dict[str, str],
) -> str:
    result = text

    for key, value in values.items():
        result = result.replace(
            "{{" + key + "}}",
            value,
        )

    return result


def _set_cell_text(
    cell,
    value: str,
) -> None:
    if not value.strip():
        return

    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        paragraph.text = value

        for extra in cell.paragraphs[1:]:
            extra.text = ""
    else:
        cell.text = value


def render_docx_template(
    template_bytes: bytes,
    doc: DocumentContent,
) -> bytes:
    document = DocxDocument(
        BytesIO(template_bytes)
    )

    field_values = {
        key.strip(): value.strip()
        for key, value in doc.metadata.items()
        if key.strip() and value.strip()
    }

    if doc.title.strip():
        field_values.setdefault(
            "문서명",
            doc.title.strip(),
        )
        field_values.setdefault(
            "제목",
            doc.title.strip(),
        )
        field_values.setdefault(
            "title",
            doc.title.strip(),
        )

    section_values = _collect_sections(doc)

    normalized_fields = {
        _normalize(key): value
        for key, value in field_values.items()
    }

    normalized_sections = {
        _normalize(key): value
        for key, value in section_values.items()
    }

    placeholder_values = {
        **field_values,
        **section_values,
    }

    for paragraph in document.paragraphs:
        replaced = _replace_placeholders(
            paragraph.text,
            placeholder_values,
        )

        if replaced != paragraph.text:
            paragraph.text = replaced

    for table in document.tables:
        for row_index, row in enumerate(
            table.rows
        ):
            for cell_index, cell in enumerate(
                row.cells
            ):
                original_text = cell.text.strip()

                replaced = _replace_placeholders(
                    original_text,
                    placeholder_values,
                )

                if replaced != original_text:
                    _set_cell_text(
                        cell,
                        replaced,
                    )
                    original_text = replaced

                normalized = _normalize(
                    original_text
                )

                field_value = normalized_fields.get(
                    normalized
                )

                if (
                    field_value
                    and cell_index + 1 < len(row.cells)
                ):
                    target = row.cells[
                        cell_index + 1
                    ]

                    if not target.text.strip():
                        _set_cell_text(
                            target,
                            field_value,
                        )

                section_value = normalized_sections.get(
                    normalized
                )

                if (
                    section_value
                    and row_index + 1 < len(table.rows)
                ):
                    target = table.rows[
                        row_index + 1
                    ].cells[cell_index]

                    if not target.text.strip():
                        _set_cell_text(
                            target,
                            section_value,
                        )

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()
